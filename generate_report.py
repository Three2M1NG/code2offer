"""生成 code2offer 数据库课程设计报告 — 完整版"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJ = os.path.dirname(__file__)
OUT = os.path.expanduser(r"~\Desktop\code2offer数据库课程设计报告.docx")
ER = os.path.join(PROJ, "er_diagram.png")

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)

# ── styles ──
s = doc.styles['Normal']; s.font.name = '宋体'; s.font.size = Pt(12)
s.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
s.paragraph_format.line_spacing = 1.5; s.paragraph_format.space_after = Pt(4)
for lv in range(1,4):
    hs = doc.styles[f'Heading {lv}']; hf = hs.font
    hf.color.rgb = RGBColor(0,0,0); hf.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.paragraph_format.first_line_indent = Cm(0)
    hf.size = Pt([15,13,12][lv-1])

def para(text, bold=False, sz=12, al=None, fi=True, fn='宋体'):
    p = doc.add_paragraph()
    if not fi: p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(sz)
    r.font.name = fn; r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if al is not None: p.alignment = al
    return p

def bullet(text, sz=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.first_line_indent = Cm(0)
    for r in p.runs: r.text = ''
    r = p.add_run(text); r.font.size = Pt(sz); r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def cell_text(cell, text, bold=False, sz=9, al=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''; pa = cell.paragraphs[0]; pa.alignment = al
    pa.paragraph_format.first_line_indent = Cm(0)
    r = pa.add_run(str(text)); r.bold = bold; r.font.size = Pt(sz)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def shdr(row, c="D9E2F3"):
    for cl in row.cells:
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), c); sh.set(qn('w:val'), 'clear')
        cl._tc.get_or_add_tcPr().append(sh)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): cell_text(t.rows[0].cells[i], h, True, 9)
    shdr(t.rows[0])
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            cell_text(t.rows[ri+1].cells[ci], val, False, 8.5,
                      WD_ALIGN_PARAGRAPH.LEFT if ci>=3 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

def kv(title, desc, sz=12):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    r1=p.add_run(f'{title}：'); r1.bold=True; r1.font.size=Pt(sz); r1.font.name='宋体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    r2=p.add_run(desc); r2.font.size=Pt(sz); r2.font.name='宋体'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

def img(path, w=5.2):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(path, width=Inches(w))
    doc.add_paragraph()

def code_block(text, sz=8):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.size = Pt(sz); r.font.name = 'Consolas'

# ═══════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════
for _ in range(5): doc.add_paragraph()
para('数据库系统基础课程设计', bold=True, sz=24, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False, fn='黑体')
doc.add_paragraph()
para('code2offer — AI 算法面试教练', bold=True, sz=18, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False, fn='黑体')
para('数据库设计与实现', bold=True, sz=14, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False, fn='黑体')
doc.add_paragraph(); doc.add_paragraph()
para('学    院：信息科学与工程学院', sz=13, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False)
para('专    业：计算机科学与技术', sz=13, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False)
para('提交日期：2026年5月', sz=13, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False)
doc.add_page_break()

# ═══════════════════════════════════════════
# 一、需求分析
# ═══════════════════════════════════════════
doc.add_heading('一、数据库需求分析', 1)

doc.add_heading('1.1 项目背景', 2)
para('code2offer 是一款面向程序员算法面试练习的 AI 教练应用。用户以文本或语音描述解题思路，系统利用大语言模型结合向量数据库进行语义检索（RAG），匹配相关 LeetCode 题目及标准题解，从正确性、复杂度分析、表述清晰度、边界情况和反应速度五个维度进行结构化评分。数据库需同时支撑结构化元数据存储和 768 维语义向量的高性能检索。')

doc.add_heading('1.2 功能需求', 2)
para('数据库需要支撑五类核心操作：')
for t,d in [
    ('题目管理','存储 15 道经典 LeetCode 高频题，覆盖数组、链表、树、DP、字符串、图等类别，支持按难度（easy/medium/hard）和标签筛选，提供题目详情查询。'),
    ('语义检索','将用户自然语言输入通过 BGE Embedding 模型转化为 768 维向量，利用 pgvector IVFFlat 索引进行余弦相似度检索，返回最匹配题目及标准解法，作为 LLM 判卷的上下文。这是系统区别于传统题库的核心能力。'),
    ('判卷记录','持久化每次判卷的完整上下文——用户输入、语音转写、AI 反馈 JSON、综合得分、匹配相似度、耗时等，支持按用户、题目、时间维度检索历史。'),
    ('评价体系管理','5 个评价维度（正确性 30%、复杂度分析 20%、表述清晰度 20%、边界情况 15%、反应速度 15%）可配置化存储，权重由触发器约束 ≤100%。'),
    ('用户管理（预留）','为登录注册功能预留用户表，当前 MVP 以匿名模式运行，但表结构已完整设计。'),
]: kv(t,d)

doc.add_heading('1.3 数据实体', 2)
for t,d in [
    ('用户 (users)','用户名、邮箱、bcrypt 密码哈希、注册与更新时间。'),
    ('题目 (problems)','中英文标题、难度（CHECK easy/medium/hard）、中文描述、标准解法与代码、时空复杂度、关键知识点、常见错误、768 维 Embedding 语义向量。'),
    ('标签 (tags)','26 个算法标签字典（Array, Hash Table, DP, Tree, BFS, DFS, Two Pointers 等），独立建表保证名称全局唯一。'),
    ('题目-标签关联 (problem_tags)','M:N 中间表，以 (problem_id, tag_id) 为复合主键，共 41 条关联记录。'),
    ('评价维度 (evaluation_criteria)','5 个维度：正确性(30%)、复杂度分析(20%)、表述清晰度(20%)、边界情况(15%)、反应速度(15%)。'),
    ('判卷记录 (analysis_history)','关联用户与题目（均可为空），存储输入文本、ASR 结果、AI 反馈 JSONB、综合得分、匹配相似度、耗时。'),
    ('评价明细 (evaluation_details)','将 JSONB 反馈拆分为结构化行级存储，每维度一条记录（得分+评语+建议），便于 SQL 聚合分析。'),
]: kv(t,d)

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、总体框架
# ═══════════════════════════════════════════
doc.add_heading('二、总体框架（ER 图）', 1)

doc.add_heading('2.1 ER 图', 2)
para('本系统共 7 张表：6 个核心实体 + 1 个 M:N 关联表。不同颜色区分实体类型：蓝色=强实体，绿色=核心事务，红色=弱实体，紫色=配置表，橙色=关联表。', fi=False)
img(ER, 5.0)
para('图 2-1  code2offer 数据库 ER 图', sz=10, al=WD_ALIGN_PARAGRAPH.CENTER, fi=False)

doc.add_heading('2.2 关系说明', 2)
for t,d in [
    ('users 1:N analysis_history','用户删除时 user_id 置 NULL（ON DELETE SET NULL），保留匿名历史。'),
    ('problems 1:N analysis_history','题目可被多次判卷，删除时级联删除关联记录。'),
    ('problems M:N tags','通过 problem_tags 复合主键实现。15 道题 × 26 个标签 = 41 条关联，平均每题 2.7 个标签。'),
    ('history 1:N evaluation_details','每次判卷产生 5 条维度明细，级联删除。'),
    ('criteria 1:N details','删除维度时 RESTRICT 防止明细数据断裂。'),
]: kv(t,d,11)

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、具体设计
# ═══════════════════════════════════════════
doc.add_heading('三、具体设计', 1)

doc.add_heading('3.1 技术选型与规范化', 2)
para('选用 PostgreSQL 16 + pgvector 扩展，支持 ACID 事务、MVCC 并发、JSONB 半结构化存储及 vector(N) 类型。设计遵循 3NF：所有非主属性完全函数依赖于主键，不存在传递依赖。problem_tags 复合主键实现 M:N 且无部分依赖。')

doc.add_heading('3.2 数据字典', 2)

para('表 3-1  users（用户表）', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','UUID','PK, DEFAULT gen_random_uuid()','用户唯一标识'],
    ['username','VARCHAR(50)','NOT NULL, UNIQUE','用户名'],
    ['email','VARCHAR(200)','NOT NULL, UNIQUE','邮箱'],
    ['password_hash','VARCHAR(255)','NOT NULL','bcrypt 哈希'],
    ['avatar_url','TEXT','NULL','头像 URL'],
    ['created_at','TIMESTAMPTZ','NOT NULL, DEFAULT now()','注册时间'],
    ['updated_at','TIMESTAMPTZ','NOT NULL','触发器自动维护'],
])

para('表 3-2  problems（题目表）—— 核心实体，不含 tags 列（通过关联表实现 M:N）', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','UUID','PK','题目唯一标识'],
    ['title','VARCHAR(200)','NOT NULL','英文标题'],
    ['title_cn','VARCHAR(200)','NULL','中文标题'],
    ['difficulty','VARCHAR(10)','NOT NULL, CHECK(IN easy,medium,hard)','难度'],
    ['description_cn','TEXT','NULL','中文题目描述'],
    ['solution_approach','TEXT','NULL','标准解题思路'],
    ['solution_code','TEXT','NULL','标准 Python 代码'],
    ['complexity_time','VARCHAR(50)','NULL','时间复杂度'],
    ['complexity_space','VARCHAR(50)','NULL','空间复杂度'],
    ['key_points','TEXT[]','NULL','关键知识点列表'],
    ['common_mistakes','TEXT[]','NULL','常见错误列表'],
    ['embedding','vector(768)','NULL','BGE-base-zh-v1.5 768维向量'],
    ['created_at','TIMESTAMPTZ','NOT NULL, DEFAULT now()','创建时间'],
    ['updated_at','TIMESTAMPTZ','NOT NULL','触发器自动维护'],
])

para('表 3-3  tags（标签字典表）—— 26 个标签，独立建表保证一致性', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','SERIAL','PK','自增 ID'],
    ['name','VARCHAR(50)','NOT NULL, UNIQUE','标签名称（如 DP, BFS）'],
    ['description','TEXT','NULL','标签说明'],
])

para('表 3-4  problem_tags（题目-标签关联表）—— 41 条关联记录', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['problem_id','UUID','PK, FK→problems CASCADE','题目外键'],
    ['tag_id','INT','PK, FK→tags CASCADE','标签外键'],
])

para('表 3-5  evaluation_criteria（评价维度配置表）', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','SERIAL','PK','自增 ID'],
    ['name','VARCHAR(50)','NOT NULL','中文名（如"正确性"）'],
    ['name_en','VARCHAR(50)','NOT NULL, UNIQUE','英文标识'],
    ['weight','NUMERIC(5,2)','CHECK(>0 AND ≤100)','权重百分比'],
    ['max_score','INT','NOT NULL, DEFAULT 10','满分'],
    ['description','TEXT','NULL','评估标准'],
    ['sort_order','INT','NOT NULL, DEFAULT 0','排序序号'],
])
para('种子数据 — 5 个评价维度：', bold=True, sz=11, fi=False)
tbl(['维度','英文标识','权重','满分','评估内容'],[
    ['正确性','correctness','30%','10','算法思路是否正确'],
    ['复杂度分析','complexity','20%','10','时间/空间复杂度分析准确性'],
    ['表述清晰度','clarity','20%','10','逻辑通顺度'],
    ['边界情况','edge_cases','15%','10','空输入、极值等边界考虑'],
    ['反应速度','delivery','15%','10','回答连贯性（模拟面试压力）'],
])

para('表 3-6  analysis_history（判卷记录表）', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','UUID','PK','记录唯一标识'],
    ['user_id','UUID','FK→users SET NULL','用户外键（可空）'],
    ['problem_id','UUID','FK→problems SET NULL','题目外键（可空）'],
    ['user_input','TEXT','NULL','文本答题内容'],
    ['asr_text','TEXT','NULL','语音转写文本'],
    ['audio_url','TEXT','NULL','音频文件路径'],
    ['ai_feedback','JSONB','NULL','AI 原始反馈 JSON'],
    ['overall_score','NUMERIC(4,2)','NULL','综合得分 (0-10)'],
    ['match_similarity','NUMERIC(5,4)','NULL','RAG 余弦相似度'],
    ['elapsed_ms','INT','NULL','耗时（毫秒）'],
    ['created_at','TIMESTAMPTZ','NOT NULL, DEFAULT now()','判卷时间'],
])

para('表 3-7  evaluation_details（评价明细表）', bold=True, sz=11, fi=False)
tbl(['字段名','类型','约束','说明'],[
    ['id','SERIAL','PK','自增 ID'],
    ['history_id','UUID','FK→history CASCADE','所属判卷记录'],
    ['criteria_id','INT','FK→criteria RESTRICT','评价维度'],
    ['score','NUMERIC(4,1)','CHECK(0≤score≤10)','该维度得分'],
    ['comment','TEXT','NULL','AI 评语'],
    ['suggestion','TEXT','NULL','改进建议'],
])

doc.add_heading('3.3 视图设计', 2)
for t,d in [
    ('v_problem_stats','按难度统计题目数、平均得分、判卷次数。LEFT JOIN 确保无判卷记录的难度也展示，CASE 保证 easy→medium→hard 顺序。'),
    ('v_user_performance','汇总每用户判卷总次数、平均/最高/最低分。'),
    ('v_recent_analysis','四表 JOIN（history×users×problems×tags），string_agg 聚合多标签为逗号分隔字符串，供前端直接消费。'),
]: kv(t,d,11)

doc.add_heading('3.4 触发器设计', 2)
para('fn_update_timestamp：BEFORE UPDATE ON users/problems，自动将 updated_at 设为当前时间。', fi=False)
para('fn_validate_criteria_weight：BEFORE INSERT/UPDATE ON evaluation_criteria，计算权重总和，超过 100% 则 RAISE EXCEPTION。', fi=False)
para('SQL 实现：')
code_block("""CREATE OR REPLACE FUNCTION fn_validate_criteria_weight()
RETURNS TRIGGER AS $$
DECLARE
    total_weight NUMERIC;
BEGIN
    SELECT COALESCE(SUM(weight), 0) INTO total_weight FROM evaluation_criteria;
    IF total_weight > 100 THEN
        RAISE EXCEPTION 'Weight sum exceeds 100%%: current=%, max=100', total_weight;
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;""", sz=7.5)

doc.add_heading('3.5 索引设计', 2)
for t,d in [
    ('idx_problems_embedding (IVFFlat, lists=10)','pgvector 近似最近邻索引。K-Means 分 10 个聚类桶，查询仅扫描最近桶，复杂度 O(√n)。余弦距离算子 <=>。'),
    ('idx_problems_difficulty (B-tree)','加速 WHERE difficulty = medium 等值筛选。'),
    ('idx_analysis_history_user / idx_analysis_history_problem','加速按用户或题目外键 JOIN 查询。'),
    ('idx_analysis_history_created (B-tree DESC)','支持 ORDER BY created_at DESC 利用索引逆序扫描。'),
    ('idx_evaluation_details_history (B-tree)','加速按 history_id 关联评价明细。'),
]: kv(t,d,11)

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、功能实现
# ═══════════════════════════════════════════
doc.add_heading('四、功能实现', 1)

doc.add_heading('4.1 环境部署', 2)
para('使用 Docker Compose 部署 pgvector/pgvector:pg16 容器。init.sql 挂载至 docker-entrypoint-initdb.d 自动建表、创建索引/视图/触发器、写入 26 个标签和 5 个评价维度的种子数据。启动命令：docker-compose up -d')

doc.add_heading('4.2 题目数据入库', 2)
para('data/seed.py 工作流程：')
for i,step in enumerate([
    '读取 data/problems.json（15 道题，含中文标签名如"数组""动态规划""双指针"等）。',
    '加载 BAAI/bge-base-zh-v1.5 模型，对每道题拼合"标题+标签+描述+题解"文本 → encode() → 768 维归一化向量。',
    '清空旧数据（DELETE FROM evaluation_details/analysis_history/problem_tags/problems）。',
    '逐条 INSERT 到 problems 表（不含 tags 列），获取 problem_id。',
    '查询中文标签→英文标签映射表，在 tags 表中查对应 tag_id，插入 problem_tags 关联记录。',
    '若标签不存在则自动 INSERT 到 tags 表。',
],1):
    para(f'({i}) {step}', sz=11, fi=False)

doc.add_heading('4.3 RAG 语义检索', 2)
para('核心代码 backend/app/services/retriever.py，三步：')
for t,d in [
    ('向量化','SentenceTransformer.encode(user_text)→768 维单位向量。模型 BAAI/bge-base-zh-v1.5。'),
    ('检索','SELECT ... FROM problems ORDER BY embedding <=> %s::vector LIMIT 3。<=> 为 pgvector 余弦距离算子，自动利用 IVFFlat 索引。'),
    ('过滤','取 Top-1，cosine_similarity=1-cosine_distance。≥0.35 匹配成功→注入 LLM Prompt；<0.35 降级为通用指导。'),
]: kv(t,d,11)

doc.add_heading('4.4 判卷与持久化', 2)
para('POST /api/v1/analyze 以 SSE 流式返回：接收 JSON {"text":"..."}→RAG 检索→prompt_manager 组装 System Prompt→DeepSeekV4Pro 流式生成（超时 fallback GLM-5.1）→正则提取 JSON 维度得分→逐 event 推送客户端→写入 analysis_history 和 evaluation_details。')
para('语音路径：Flutter 录音(.m4a)→POST /api/v1/transcribe→Qwen3-ASR-Flash 转写→文本填入输入框→用户编辑→/api/v1/analyze。')

doc.add_heading('4.5 系统验证', 2)
para('以下为数据库重建后的全链路验证数据（2026-05-22 实际运行结果）：', bold=True, fi=False)

para('(1) 容器与表结构验证：', bold=True, sz=11, fi=False)
code_block("""$ docker ps
NAMES             STATUS         PORTS
code2offer-db-1   Up (healthy)   0.0.0.0:5432->5432/tcp

$ psql -c "SELECT table_name FROM information_schema.tables WHERE table_schema='public'"
 analysis_history
 evaluation_criteria
 evaluation_details
 problem_tags
 problems
 tags
 users
(7 rows)""", sz=7.5)

para('(2) 数据完整性验证：', bold=True, sz=11, fi=False)
code_block("""problems:         15 rows   (15 道算法题全部入库)
tags:             26 rows   (Array, DP, BFS, DFS, Two Pointers, Union Find ...)
problem_tags:     41 rows   (15 题与 26 标签的 M:N 关联)
evaluation_criteria: 5 rows (正确性/复杂度/清晰度/边界/反应速度，权重合计 100%)

-- 孤立题目检查（全部关联成功）
SELECT p.title FROM problems p
LEFT JOIN problem_tags pt ON p.id = pt.problem_id
WHERE pt.problem_id IS NULL;
→ (0 rows)  // 无孤立题目

-- 按标签筛选验证
SELECT p.title FROM problems p
JOIN problem_tags pt ON p.id = pt.problem_id
JOIN tags t ON t.id = pt.tag_id
WHERE t.name = 'DP';
→ 6 problems: Best Time to Buy and Sell Stock, Maximum Subarray,
  Longest Palindromic Substring, Coin Change, Climbing Stairs, Word Break""", sz=7.5)

para('(3) API 端点验证：', bold=True, sz=11, fi=False)
code_block("""GET /api/v1/health                    → 200 {"status":"ok","version":"0.1.0"}
GET /api/v1/problems                   → 200 15 problems, tags via M:N
GET /api/v1/problems?difficulty=medium → 200 9 problems
GET /api/v1/problems?tag=DP            → 200 6 problems (correct)
GET /api/v1/problems/{id}              → 200 full detail with tags array""", sz=7.5)

para('验证结论：M:N 架构一致性通过，所有数据完整性约束生效，API 全端点正常。', bold=True, fi=False)

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、TRIZ
# ═══════════════════════════════════════════
doc.add_heading('五、TRIZ 创新理论应用', 1)
para('TRIZ（发明问题解决理论）由苏联发明家阿奇舒勒于 1946 年创立，核心方法：识别技术矛盾 → 应用 40 个发明原理化解矛盾。本设计在数据库架构中多次运用 TRIZ 思维：')

for t,d in [
    ('技术矛盾','传统题库只能关键词精确匹配，无法"理解"用户自然语言描述的解题思路。矛盾：改善"语义理解精度"导致"系统复杂性"增加。'),
    ('原理 17 — 空间维数变化','将题目从离散标签空间映射到 768 维语义空间。在高维空间中，"Two Sum"与"用哈希表存储已遍历元素"余弦相似度 >0.85，传统 WHERE tags LIKE 无法实现。IVFFlat 索引保证高维检索性能。'),
    ('原理 35 — 参数变化','将匹配结果从布尔值变为 0-1 连续余弦相似度，支持灵活阈值策略：≥0.35 匹配题目 → 0.2-0.35 提示明确 → <0.2 降级通用指导。'),
    ('原理 24 — 中介物','引入语义向量作为用户输入与题目数据的中介层，解耦"表达方式"和"存储格式"，使口语化、中英混杂等多样表述均有鲁棒性。'),
    ('原理 1 — 分割','将 JSONB 聚合反馈拆分为 evaluation_details 结构化行，使 SQL 可直接对单维度聚合查询（如"正确性维度平均分"），无需应用层解析 JSON。'),
]: kv(t,d)

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、个人心得
# ═══════════════════════════════════════════
doc.add_heading('六、个人心得', 1)

for t,d in [
    ('理论与实践的结合','最初 tags 以 TEXT[] 存于 problems 表，经 3NF 分析后抽出 tags 字典和 problem_tags 关联表，避免了拼写不一致和反向查询困难。26 个标签、41 条关联记录的验证数据证明了规范化设计的正确性。数据库重建过程中 init.sql 与 seed.py 的"精神分裂"问题（一个用 M:N、一个假设 TEXT[] 列）更是让我切身体会到代码与 DDL 一致性的重要性。'),
    ('向量数据库的新认知','pgvector 让我认识到数据库不仅是精确匹配的机器，而是能在高维空间中度量语义距离。IVFFlat 以 K-Means 聚类 + 近似搜索替代全表扫描的设计思路，将 O(n) 降至 O(√n)，是索引本质（空间换时间）在向量领域的自然延伸。'),
    ('工程细节决定质量','ON DELETE SET NULL vs CASCADE vs RESTRICT、BEFORE vs AFTER 触发器、TIMESTAMPTZ vs TIMESTAMP、psycopg v3 vs psycopg2 驱动前缀（postgresql+psycopg://）——每个看似微小的选择都直接影响数据一致性和跨时区行为。这些无法从课本中完全领会，必须在实践中积累。'),
    ('TRIZ 方法论的启发','从"罗列功能"转变为"分析矛盾→寻找原理"，理解了为什么用向量数据库——它解决的是精确匹配与语义理解之间的根本矛盾。这种思维框架对分析任何复杂工程问题都有长远价值。'),
]: kv(t,d)

# ── save ──
doc.save(OUT)
total = sum(len(p.text) for p in doc.paragraphs)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)/1024:.1f} KB, ~{total} chars')
